#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera docs/informe_final.docx — informe técnico completo del Sistema Multiagente
Everywhere Travel, siguiendo el índice académico exacto.

El contenido se consolida desde los docs/*.md del repositorio. Los diagramas se
insertan como placeholders (recuadros) con la ruta del archivo fuente a exportar.

Uso:
    python scripts/build_report.py
Requiere: python-docx  (pip install python-docx)
"""
from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

ROOT = Path(__file__).resolve().parent.parent

doc = Document()

# ─── Estilos base ─────────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

ACCENT = RGBColor(0x1F, 0x4E, 0x79)


# ─── Helpers ──────────────────────────────────────────────────────────────────
def _shade(p, fill="F2F2F2"):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def h1(text):
    doc.add_page_break()
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = ACCENT


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = ACCENT


def h3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = ACCENT


def para(text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0] + ": ")
            r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def numbered(items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    try:
        t.style = "Light Grid Accent 1"
    except Exception:
        t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    doc.add_paragraph()
    return t


def code(text):
    p = doc.add_paragraph()
    _shade(p)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        if i < len(lines) - 1:
            r.add_break()
    return p


def placeholder(caption, source):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    _shade(cell.paragraphs[0], "EAF1F8")
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p1.add_run("[ DIAGRAMA — insertar imagen aquí ]")
    r.bold = True
    r.font.size = Pt(10)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(caption)
    r2.italic = True
    r2.font.size = Pt(9)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Fuente: " + source)
    r3.font.size = Pt(8)
    r3.font.name = "Consolas"
    doc.add_paragraph()


def add_toc():
    p = doc.add_paragraph()
    r = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    r._r.append(b); r._r.append(instr); r._r.append(sep)
    p.add_run("»» Actualice este índice en Word: seleccione todo (Ctrl+A) y presione F9, "
              "o clic derecho sobre el índice → Actualizar campos → Actualizar toda la tabla.")
    r3 = p.add_run()
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    r3._r.append(e)


def read_prompt(rel):
    fp = ROOT / rel
    try:
        return fp.read_text(encoding="utf-8").strip()
    except Exception:
        return "(no se pudo leer " + rel + ")"


# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for _ in range(4):
    doc.add_paragraph()
tr = title.add_run("Sistema Multiagente Everywhere Travel")
tr.bold = True
tr.font.size = Pt(26)
tr.font.color.rgb = ACCENT

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Informe Técnico — Automatización Inteligente de Procesos con Agentes de IA")
sr.font.size = Pt(14); sr.italic = True

for _ in range(6):
    doc.add_paragraph()

meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Universidad Privada Antenor Orrego (UPAO)\n").bold = True
meta.add_run("7.º Ciclo — Curso de Automatización Inteligente de Procesos\n")
meta.add_run("Autor: Gerson Portal\n")
meta.add_run("Versión 1.0 — 17 de julio de 2026")


# ══════════════════════════════════════════════════════════════════════════════
# CONTROL DE VERSIONES
# ══════════════════════════════════════════════════════════════════════════════
h1("Control de versiones")
table(
    ["Versión", "Fecha", "Autor", "Descripción del cambio"],
    [
        ["0.1", "2026-05-22", "Gerson Portal",
         "Arquitectura inicial: 9 agentes, Saga + RabbitMQ, MCP Envelope, tests, demo E2E; "
         "login JWT; ItineraryAgent + generación de PDF"],
        ["0.2", "2026-06-29", "Gerson Portal",
         "Mejoras en agentes backend, API e infraestructura (migración a Ollama local vía swarms_compat)"],
        ["0.3", "2026-06-30", "Gerson Portal",
         "Mejoras de UI del dashboard y componentes reutilizables"],
        ["1.0", "2026-07-17", "Gerson Portal",
         "Cierre del informe: RAG (pgvector), salida estructurada forzada, HITL con confidence real, "
         "observabilidad conectada, seguridad (JWT verificado), golden set + evaluación local, CI, "
         "alertas, dashboard Grafana, 12 ADR, BPMN AS-IS/TO-BE, catálogos de prompts/tools, ROI"],
    ],
)


# ══════════════════════════════════════════════════════════════════════════════
# TABLA DE CONTENIDOS
# ══════════════════════════════════════════════════════════════════════════════
h1("Tabla de contenidos")
add_toc()


# ══════════════════════════════════════════════════════════════════════════════
# 1. RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Resumen ejecutivo")

h2("1.1 Problema")
para("Everywhere Travel es una agencia de viajes con múltiples sedes que opera con hojas de "
     "cálculo compartidas para cotizar paquetes, coordinar reservas entre sedes y liquidar "
     "comisiones. Este modelo produce cotizaciones lentas (2-4 horas por consulta compleja), "
     "riesgo de doble reserva por falta de coordinación en tiempo real, errores de redondeo en "
     "cálculos financieros hechos a mano, y ausencia de un rastro de auditoría verificable para "
     "compliance regulatorio (IGV).")

h2("1.2 Solución propuesta")
para("Una plataforma interna operada por 9 agentes especializados (Orchestrator, Sales, "
     "Quotation, Reservation, Finance, Document, Validation, Monitoring, Notification, más "
     "Itinerary) que colaboran a través de un contrato de mensajería explícito (MCP Envelope) "
     "sobre un event bus (RabbitMQ), coordinados con el patrón Saga para garantizar que ninguna "
     "transacción distribuida quede en estado inconsistente.")
para("Un LLM local (Ollama, qwen3:8b) se usa solo donde hay ambigüedad de lenguaje natural o "
     "generación creativa genuina (interpretar la consulta del cliente, estimar paquetes "
     "personalizados, redactar itinerarios, evaluar conflictos operativos). El resto del sistema "
     "(compliance, aritmética financiera, locking, circuit breaking) es deliberadamente "
     "determinístico, con las salidas del LLM forzadas a un esquema JSON y validadas con Pydantic "
     "antes de usarse en cualquier flujo de negocio.")

h2("1.3 Resultado esperado")
bullets([
    "Cotización de paquetes de catálogo en menos de 30 segundos; personalizados en minutos.",
    "Cero doble reservas mediante locking optimista sobre Redis.",
    "Ledger financiero inmutable y auditable, con IGV y márgenes en Decimal exacto.",
    "Continuidad operativa ante fallos parciales: circuit breaker, dead-letter queue y "
    "escalación automática a un humano cuando la confianza de una decisión del LLM cae bajo umbral.",
    "Observabilidad real: métricas Prometheus con alertas, dashboard Grafana provisionado, "
    "y evaluación local (golden set) que sustituye a LangSmith sin cuenta externa.",
])


# ══════════════════════════════════════════════════════════════════════════════
# 2. ANÁLISIS
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Análisis")

h2("2.1 Justificación: ¿se necesita un LLM?")
para("Sí, para un subconjunto acotado de tareas; el resto del sistema es deliberadamente "
     "determinístico. Un LLM aporta valor real donde hay ambigüedad de lenguaje natural o "
     "generación creativa que no se puede resolver con reglas fijas:")
table(
    ["Tarea", "¿Por qué LLM y no reglas?"],
    [
        ["Interpretar consulta y armar PackageRequest (SalesAgent)",
         "Las preferencias llegan en texto libre ('hotel 4*', 'vuelo incluido'); mapearlas a una "
         "estructura requiere comprensión de lenguaje."],
        ["Estimar componentes de paquete personalizado (QuotationAgent)",
         "Sin paquete de catálogo, componer un desglose razonable para un destino arbitrario "
         "requiere razonamiento."],
        ["Redactar itinerario día a día (ItineraryAgent)",
         "Generación de contenido creativo en español con contexto cultural."],
        ["Evaluar conflicto operativo y decidir escalación (Orchestrator)",
         "Juzgar si un conflicto es un problema de integridad, con qué confianza, no es una regla "
         "determinista."],
    ],
)
para("Todo lo demás —cálculo de IGV, cronogramas de pago, motor de reglas de compliance "
     "(R001-R012), circuit breaker, deduplicación, generación de PDF— es intencionalmente "
     "determinístico. Usar un LLM para compliance regulatorio o aritmética financiera "
     "introduciría no-determinismo donde la rúbrica de negocio exige exactitud auditable.")

h2("2.2 Objetivos y alcance")
h3("Objetivo general")
para("Automatizar el ciclo de vida completo de una venta de paquete turístico (consulta → "
     "cotización → validación → reserva → liquidación → documentos → itinerario) mediante agentes "
     "especializados que colaboran vía un contrato de mensajería explícito (MCP), sin que ningún "
     "agente concentre toda la lógica de negocio.")
h3("Objetivos específicos")
numbered([
    "Eliminar la doble reserva y los errores de cálculo manual mediante agentes deterministas "
    "para dinero y disponibilidad.",
    "Reducir el tiempo de cotización de horas a segundos mediante LLM acotado + catálogo.",
    "Dar trazabilidad completa de cada transacción distribuida (patrón Saga + audit log inmutable).",
    "Sostener continuidad operativa ante fallos parciales (circuit breaker, dead-letter, escalación humana).",
])
h3("Dentro del alcance")
para("Cotización, validación, reserva, liquidación, documentos, itinerarios, RAG sobre "
     "catálogo/destinos, observabilidad (métricas + logs estructurados + golden set local).")
h3("Fuera del alcance")
para("Pagos reales (pasarela de pago), integración con GDS/aerolíneas reales, multi-tenant "
     "(multi-agencia), app móvil nativa, LangSmith/tracing en la nube.")

h2("2.3 Requisitos funcionales")
bullets([
    ("RF-01", "Generar una cotización a partir de una consulta en < 30s (catálogo) o con estimación LLM (personalizado)."),
    ("RF-02", "Toda cotización debe pasar por ValidationAgent antes de poder confirmarse."),
    ("RF-03", "Una reserva debe fallar de forma segura (sin doble-booking) ante solicitudes concurrentes."),
    ("RF-04", "Toda liquidación completa debe disparar generación automática de factura y comprobante."),
    ("RF-05", "El sistema debe generar un itinerario descargable en PDF por cotización validada."),
    ("RF-06", "Un conflicto entre agentes debe resolverse automáticamente o escalarse a un humano según el nivel de confianza."),
])

h2("2.4 Requisitos no funcionales (propios de IA)")
bullets([
    ("RNF-01 (Confiabilidad de salida)", "Toda respuesta del LLM debe validar contra un JSON Schema forzado; si falla, se usa un fallback determinístico."),
    ("RNF-02 (Costo)", "El proveedor LLM por defecto debe tener costo marginal cero por ejecución (Ollama local)."),
    ("RNF-03 (Latencia)", "p95 de llamadas LLM reportado en Prometheus, con alerta si supera 20s."),
    ("RNF-04 (Trazabilidad)", "Toda llamada LLM debe registrarse en agent_interaction_logs (input, output, duración, éxito)."),
    ("RNF-05 (Degradación con gracia)", "Si el LLM no está disponible, el flujo de negocio debe completarse con lógica determinística."),
])

h2("2.5 Inventario de conocimiento y acciones")
h3("Fuentes de conocimiento (para RAG)")
table(
    ["Fuente", "Contenido", "Tabla"],
    [
        ["Catálogo de paquetes", "Nombre, destino, precio, duración, incluye/excluye", "packages (embedding vector(768))"],
        ["Guías de destino curadas", "Clima, altitud, cultura, tips por destino (core/rag/content.py)", "destination_knowledge"],
    ],
)
h3("Acciones externas (tools)")
table(
    ["Tool", "Agente", "Efecto"],
    [
        ["_tool_select_package / _tool_validate_dates / _tool_build_customizations", "Sales", "Cómputo local, sin I/O externo"],
        ["_tool_semantic_search_packages", "Sales", "Llama a GET /packages/semantic-search (RAG)"],
        ["_tool_calculate_igv / _tool_check_margin_policy / _tool_estimate_component_price / _tool_detect_budget_anomaly", "Quotation", "Cómputo local"],
        ["_tool_get_destination_info (vía RAG)", "Itinerary", "Llama a GET /knowledge/destinations/search, con fallback estático"],
        ["_tool_calculate_days / _tool_get_included_services", "Itinerary", "Cómputo local"],
    ],
)
para("La ficha completa de cada tool (Args, Returns, I/O, errores) está en la sección 3.4.")

h2("2.6 Criterios de éxito y conjunto de evaluación (golden set)")
para("El golden set (tests/evaluation/golden_set.py) contiene 12 casos que verifican que la capa "
     "de salida estructurada (a) acepta JSON válido —limpio o envuelto en prosa/bloques de "
     "código— y (b) rechaza salidas incompletas o mal tipadas, forzando el fallback determinístico. "
     "Criterio de éxito: 100% de los casos deben pasar en cada PR (ver .github/workflows/ci.yml). "
     "Procedimiento y reporte completos en la sección 5.")

h2("2.7 Análisis de riesgos")
table(
    ["Riesgo", "Prob.", "Impacto", "Mitigación"],
    [
        ["El LLM local (8B) alucina o produce JSON inválido", "Media", "Medio", "Salida forzada por schema + fallback determinístico (ADR-010)"],
        ["Doble reserva por condición de carrera", "Baja", "Alto", "Optimistic locking Redis SETNX (ADR-004)"],
        ["Un agente cae y bloquea la Saga", "Media", "Alto", "Dead-letter queue + reintentos + escalación humana tras 3 fallos"],
        ["Ollama no disponible", "Media", "Bajo", "Cada agente LLM tiene fallback determinístico; el sistema sigue vendiendo"],
        ["Routers de la API sin autenticación real", "Alta (mitigada)", "Alto", "get_current_user verifica JWT; auth servicio-a-servicio pendiente"],
        ["JWT accesible desde JS (XSS)", "Media", "Medio", "Cookie httpOnly añadida como defensa adicional"],
        ["pgvector con catálogo pequeño", "Baja", "Bajo", "Aceptado; el volumen no justifica índice ANN afinado (ADR-009)"],
    ],
)


# ══════════════════════════════════════════════════════════════════════════════
# 3. DISEÑO
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Diseño")

h2("3.1 Arquitectura general")
h3("Patrón de orquestación elegido")
para("Orquestación mediante el patrón Saga (core/saga_coordinator.py) sobre un event bus "
     "(RabbitMQ, topic exchange + dead-letter), con cada agente como proceso/contenedor "
     "independiente. No se usa LangGraph; la justificación está en el ADR-001. Cada paso de una "
     "Saga se registra en Redis (hot) y se sincroniza a PostgreSQL (cold, audit trail permanente). "
     "Los agentes se comunican exclusivamente vía mensajes MCPEnvelope enrutados por topic "
     "exchange; no hay llamadas RPC directas entre agentes.")
para("Topología híbrida jerárquica-estrella (ADR-003): el OrchestratorAgent centraliza el "
     "enrutamiento inicial y la gestión de Sagas, pero dentro de cada dominio los agentes se "
     "comunican directamente (Sales → Quotation → Validation) sin pasar por el Orchestrator en "
     "cada hop.")
h3("Esquema de la composición")
placeholder("Diagrama de componentes: capas Frontend, API Gateway, Orquestación, Agentes, "
            "Infraestructura (PostgreSQL/Redis/RabbitMQ/MinIO).",
            "docs/architecture.md — sección 1 (Mermaid)")

h2("3.2 Diagrama de proceso (BPMN)")
para("Se modelan dos procesos completos, de la consulta del cliente al cierre de la venta:")
bullets([
    ("AS-IS (proceso manual, antes del ERP)", "5 lanes (Cliente, Vendedor, Gerente de sede, "
     "Operaciones, Finanzas); muestra los cuellos de botella: aprobación del gerente, loop de "
     "cambios que rehace la cotización, y la rama de doble reserva resuelta por teléfono."),
    ("TO-BE (con el ERP multiagente)", "9 lanes, uno por agente; las decisiones manuales pasan a "
     "gateways automáticos (¿regla BLOCKING? con rama de escalación HITL, ¿lock adquirido? con "
     "resolución automática de conflicto, ¿saldo = 0?)."),
])
placeholder("BPMN AS-IS — proceso comercial manual (5 lanes).",
            "docs/bpmn/as_is_proceso_manual.bpmn (abrir en demo.bpmn.io y exportar imagen)")
placeholder("BPMN TO-BE — proceso comercial con el ERP multiagente (9 lanes).",
            "docs/bpmn/to_be_proceso_erp.bpmn (abrir en demo.bpmn.io y exportar imagen)")
para("La narrativa detallada, fase por fase, de ambos procesos está en docs/proceso_negocio.md.")

h2("3.3 Subsistema RAG")
para("SalesAgent e ItineraryAgent recuperan conocimiento por similaridad semántica en vez de "
     "match exacto de string (justificación en ADR-009). Se usa pgvector (extensión de "
     "PostgreSQL) para almacenar embeddings de paquetes (packages.embedding) y de una base de "
     "conocimiento curada de destinos (destination_knowledge), generados con el modelo de "
     "embeddings local de Ollama (nomic-embed-text, 768 dimensiones). La recuperación (similaridad "
     "de coseno) se expone vía endpoints propios que los agentes consultan por HTTP.")
para("Flujo de recuperación: la consulta del cliente se embebe con el mismo modelo, y "
     "'ORDER BY embedding <=> :query_embedding LIMIT k' (operador de distancia de coseno de "
     "pgvector) devuelve los k resultados más cercanos. Si Ollama no está disponible, el endpoint "
     "responde 503 y el agente cae a su fallback determinístico.")
placeholder("Subsistema RAG: fuentes → indexación (embeddings Ollama) → pgvector → consumo (Sales, Itinerary).",
            "docs/architecture.md — sección 7bis (Mermaid)")

h2("3.4 Especificación de herramientas (tools)")
para("Todas las tools son funciones Python síncronas (restricción de swarms.Agent), registradas "
     "en el tools=[...] del agente. Con Ollama, el modelo no las invoca vía function-calling "
     "nativo: su nombre y primera línea de docstring se inyectan como contexto en el prompt. A "
     "continuación, la ficha de cada herramienta.")

_tools = [
    ("_tool_select_package", "SalesAgent", "Selecciona el mejor paquete de catálogo dado presupuesto y destino",
     "packages_json:str, budget_max:float, destination:str", "JSON {selected, name, price} o {selected:null, reason}",
     "Ninguno (cómputo puro)"),
    ("_tool_validate_dates", "SalesAgent", "Valida que las fechas sean lógicas y con ≥48h de anticipación",
     "start_date:str, end_date:str", "JSON {valid, duration_days, advance_days, issues[]}", "Ninguno"),
    ("_tool_build_customizations", "SalesAgent", "Construye personalizaciones desde preferencias en texto libre",
     "preferences_list:str, budget_min:float, budget_max:float", "JSON {hotel_category, includes_flight, ...}", "Ninguno"),
    ("_tool_semantic_search_packages", "SalesAgent", "Búsqueda RAG de paquetes por similaridad semántica (fallback)",
     "query:str, top_k:int=5", "JSON {packages:[...]}", "GET /packages/semantic-search (HTTP); 503 → {packages:[]}"),
    ("_tool_calculate_igv", "QuotationAgent", "Calcula IGV (18%) sobre un monto base",
     "base_amount:float", "JSON {igv_amount, total_with_igv, rate_used}", "No falla"),
    ("_tool_check_margin_policy", "QuotationAgent", "Valida si un margen cumple la política mínima (15%)",
     "margin_pct:float", "JSON {compliant, minimum_required, recommendation, severity}", "No falla"),
    ("_tool_estimate_component_price", "QuotationAgent", "Estima el precio de un componente según destino y duración",
     "component_type:str, destination:str, traveler_count:int, duration_days:int", "JSON {concept, unit_price, quantity, subtotal}", "No falla (precio por defecto)"),
    ("_tool_detect_budget_anomaly", "QuotationAgent", "Detecta anomalías: sobre-presupuesto, margen bajo, costo cero",
     "total_cost:float, budget_max:float, margin_pct:float", "JSON {anomaly_flags[], has_anomalies, severity}", "No falla"),
    ("_tool_get_destination_info", "ItineraryAgent", "Info de destino: RAG primero, diccionario estático como fallback",
     "destination:str", "JSON {destination, source, knowledge[] o climate/...}", "GET /knowledge/destinations/search; fallback estático"),
    ("_tool_calculate_days", "ItineraryAgent", "Calcula distribución de días del viaje",
     "start_date:str, end_date:str", "JSON {total_days, arrival_day, departure_day, ...}", "Fallback 5 días"),
    ("_tool_get_included_services", "ItineraryAgent", "Extrae servicios incluidos desde line_items",
     "line_items_json:str", "JSON {included[], not_included[]}", "Fallback lista genérica"),
]
for name, agent, prop, args, ret, err in _tools:
    h3("Ficha de herramienta — " + name)
    table(
        ["Campo", "Valor"],
        [["Agente", agent], ["Propósito", prop], ["Args", args], ["Returns", ret], ["I/O y errores", err]],
    )
para("Todas las tools son puras o de solo-lectura: ninguna escribe en PostgreSQL, Redis ni "
     "RabbitMQ directamente. La persistencia de la decisión (guardar cotización, publicar evento "
     "MCP) siempre la ejecuta el código determinístico del agente, nunca la tool.")

h2("3.5 Orquestación con estado (LangGraph)")
para("NOTA: el sistema NO utiliza LangGraph. Se documenta aquí el equivalente funcional que sí "
     "implementa —orquestación con estado mediante Saga + event bus— mapeando cada concepto de "
     "LangGraph a su realización concreta en este sistema (justificación de la elección en ADR-001).")
table(
    ["Concepto (LangGraph)", "Equivalente real en el sistema"],
    [
        ["Estado compartido (TypedDict / Pydantic)",
         "MCPEnvelope y payloads tipados con Pydantic v2 (core/mcp/envelope.py) + estado de Saga en Redis (saga:{id})"],
        ["Nodos del grafo",
         "Los 9 agentes especializados (agents/*/agent.py), cada uno como proceso/contenedor independiente"],
        ["Aristas y condicionales",
         "ROUTING_TABLE del Orchestrator + lógica if/else (p. ej. chequeo de circuit breaker, gateways de conflicto)"],
        ["Persistencia y checkpointing",
         "Log de pasos de la Saga en Redis (hot, TTL 1h) sincronizado a PostgreSQL (cold, audit trail). "
         "No hay checkpointer de grafo; la reanudación es por detección de sagas estancadas (MonitoringAgent)"],
    ],
)
h3("Human-in-the-loop (HITL)")
para("El punto de entrada de HITL es la evaluación de conflictos del Orchestrator "
     "(Fase 3). Dos sub-agentes evalúan el mismo conflicto con salida forzada a JSON Schema: "
     "conflict-validation-agent (¿es problema de integridad?, con confidence 0-1) y "
     "conflict-monitoring-agent (¿impacto operativo?, needs_escalation). Si confidence < 0.7 o "
     "needs_escalation es verdadero, el MonitoringAgent ejecuta la escalación real "
     "(REQUIRES_MANUAL_INTERVENTION vía Redis pub/sub → WebSocket → dashboard). La constante "
     "HUMAN_ESCALATION_CONFIDENCE_THRESHOLD = 0.7 respalda la instrucción del prompt del "
     "Orchestrator, que antes no tenía cálculo real detrás.")
para("Otros dos puntos de escalación (deterministas): tras 3 reintentos fallidos de un mensaje "
     "dead-letter, y tras 3 fallos consecutivos de generación de documento. Limitación conocida: "
     "la escalación es de notificación, no de aprobación bloqueante; la Saga queda en "
     "REQUIRES_MANUAL esperando intervención manual directa.")

h2("3.6 Deep Agents — patrón de planificación")
para("NOTA: el sistema NO usa el patrón Deep Agent (planificador dinámico + sub-agentes ad hoc). "
     "Se documenta la decisión (ADR-012).")
h3("Cuándo elegir Deep Agent frente a un agente simple")
para("El patrón Deep Agent aporta valor cuando la tarea es abierta y de largo horizonte, y ni el "
     "número de pasos ni su orden se conocen de antemano (p. ej. 'investiga X y escribe un "
     "informe'). Everywhere Travel tiene procesos de negocio fijos y conocidos: una cotización "
     "siempre pasa por Sales → Quotation → Validation; el orden lo define una regla de negocio y "
     "debe ser predecible y auditable (compliance, IGV). Por eso se eligieron agentes "
     "especializados de flujo fijo en vez de un planificador dinámico.")
h3("Componentes a especificar / Catálogo de sub-agentes")
para("En el equivalente implementado, los 'sub-agentes' son los 9 agentes especializados con "
     "responsabilidad fija (ver sección 3.1 y docs/agent_contracts.md). El único punto con "
     "múltiples sub-agentes coordinados dinámicamente es la Fase 3 del Orchestrator "
     "(conflict-validation-agent + conflict-monitoring-agent), descrita en 3.5.")
h3("Límites operativos")
table(
    ["Agente", "max_loops", "Timeout", "Memoria", "Fallback si el LLM falla"],
    [
        ["SalesAgent", "1", "120s", "memory_chunk_size=2000", "_fallback_package_request (selección determinística)"],
        ["QuotationAgent", "2", "120s", "memory_chunk_size=1500", "_budget_fallback (75% del presupuesto máx.)"],
        ["ItineraryAgent", "1", "120s", "—", "_fallback_itinerary (plantilla por días)"],
        ["Orchestrator Fase 3 (×2)", "1", "120s", "—", "ConflictAgent simple; si falla, needs_escalation=true"],
    ],
)
para("Límites transversales: cada mensaje MCP lleva ttl_seconds=300 y retry_count≤10; "
     "MonitoringAgent descarta y escala tras 3 reintentos. Ninguna llamada LLM bloquea el event "
     "loop (asyncio.to_thread).")

h2("3.7 Esquemas de salida estructurada")
para("Cada agente con LLM define su contrato de salida como modelo Pydantic (PackageRequest, "
     "LineItemsOutput, ItineraryOutput, ConflictValidationOutput, ConflictMonitoringOutput). Su "
     ".model_json_schema() se pasa como response_schema al Agent de swarms_compat, que lo reenvía "
     "a Ollama como 'format' — el modelo queda restringido a emitir JSON conforme al schema "
     "durante el decoding (constrained decoding, Ollama ≥ 0.5), no solo 'instruido' por prompt. "
     "core/structured_output.py::parse_structured_output() valida el resultado contra el mismo "
     "modelo Pydantic como segunda capa, con extracción manual de bloques JSON como red de "
     "seguridad final antes de caer al fallback determinístico (ADR-010).")

h2("3.8 Robustez operativa")
bullets([
    ("Circuit breaker", "core/circuit_breaker.py — CLOSED → OPEN (5 fallos en 60s) → HALF_OPEN (30s cooldown) → CLOSED. Estado persistido en Redis."),
    ("Patrón Saga con compensación", "core/saga_coordinator.py — log de pasos, detección de sagas estancadas (>5 min), estados RUNNING/COMPLETED/COMPENSATING/FAILED/REQUIRES_MANUAL."),
    ("Dead-letter queue", "MonitoringAgent reencola con backoff exponencial (1s…32s); descarta y escala a humano tras 3 reintentos."),
    ("Deduplicación", "Redis processed:{message_id} (TTL 24h) evita procesar dos veces el mismo mensaje."),
    ("Degradación con gracia", "Ante fallo del LLM, cada agente completa el flujo con lógica determinística (_fallback_*)."),
])

h2("3.9 Seguridad y privacidad")
bullets([
    ("Autenticación", "JWT (HS256, PyJWT), expira a las 8h. Hallazgo corregido: antes ningún endpoint verificaba el token; ahora get_current_user protege clients, stats, itinerary. Cookie httpOnly añadida como defensa contra XSS."),
    ("Auth servicio-a-servicio (pendiente)", "packages, quotations, reservations, etc. reciben tráfico interno de agentes sin JWT de usuario; protegerlos requiere una API key interna o red de confianza."),
    ("Gestión de secretos", ".env / .env.production no versionados; SECRET_KEY con valor por defecto que debe rotarse en producción; sin vault (aceptable para alcance académico)."),
    ("Validación de entradas", "Doble capa Pydantic + JSON Schema Draft-07 en cada hop; MCPEnvelope rechaza agentes fuera de la lista blanca; tests adversariales (payload gigante, SQL injection, rangos)."),
    ("Prompt injection", "Riesgo acotado: la salida del LLM está forzada a schema y SalesAgent nunca calcula precios ni reserva; toda cotización pasa por ValidationAgent. Sin filtro explícito de inyección todavía."),
    ("Privacidad (PII)", "clients.preferences/document_number son PII sin cifrado a nivel de columna; los prompts se procesan localmente vía Ollama (no salen a un proveedor externo)."),
    ("Auditoría", "validation_logs inmutable; agent_interaction_logs registra cada interacción LLM."),
])


# ══════════════════════════════════════════════════════════════════════════════
# 4. ADR
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Registro de decisiones de arquitectura (ADR)")
para("Formato MADR simplificado (Contexto / Decisión / Alternativas / Consecuencias). Cada ADR "
     "es inmutable una vez aceptado. A continuación, los 12 ADR del proyecto.")

_adrs = [
    ("ADR-001", "Orquestación Saga + Event Bus en vez de LangGraph",
     "Coordinar 9 agentes en flujos de larga duración que sobrevivan al reinicio de cualquier agente y permitan compensación. LangGraph asume orquestación centralizada en un proceso; los agentes son contenedores independientes.",
     "Usar el patrón Saga sobre RabbitMQ (topic exchange + dead-letter). Cada paso se registra en Redis + PostgreSQL. Comunicación exclusiva vía MCPEnvelope.",
     "Cada agente escala independiente; el log de la Saga es auditoría nativa; un agente caído no bloquea a los demás. Costo: no hay visualización de grafo out-of-the-box; routing condicional en if/else."),
    ("ADR-002", "Ollama local como proveedor LLM por defecto",
     "Los agentes necesitan un LLM para tareas acotadas; el proyecto corre cientos de sagas de prueba/demo. Reemplaza la documentación previa que afirmaba usar Anthropic (el código nunca lo invocó).",
     "Ollama local (qwen3:8b), invocado por HTTP desde swarms_compat. Punto de extensión: LLM_PROVIDER=anthropic + ANTHROPIC_API_KEY activa Claude vía swarms/litellm.",
     "Costo marginal cero; sin rate limits externos; privacidad (los datos no salen). Mitigado: modelo de 8B menos confiable → schema forzado + fallback determinístico; tool-calling no nativo con Ollama."),
    ("ADR-003", "Topología híbrida jerárquica-estrella",
     "Con 9 agentes, la topología determina trazabilidad y acoplamiento.",
     "Orchestrator centraliza routing y Sagas (jerárquico); dentro de cada dominio los agentes se comunican directamente (estrella interna).",
     "Visibilidad completa sin cuello de botella de tráfico; los dominios evolucionan su comunicación interna sin tocar el Orchestrator."),
    ("ADR-004", "Optimistic locking con Redis SETNX",
     "ReservationAgent debe evitar doble reserva ante solicitudes concurrentes.",
     "Redis SETNX (lock:{type}:{id}, TTL 30s) antes de insertar, en vez de locks de BD.",
     "O(1), no bloquea el motor transaccional; TTL evita locks huérfanos. Es advisory: protege solo contra agentes que respetan el protocolo (ReservationAgent es el único escritor)."),
    ("ADR-005", "RabbitMQ topic exchange en vez de Kafka",
     "Event bus para enrutar MCPEnvelope con dead-letter, prioridad y routing por patrón.",
     "RabbitMQ topic exchange + dead-letter exchange.",
     "Routing por topic mapea a los payload_type del contrato MCP; dead-letter nativo. Costo: sin replay largo como Kafka (la auditoría vive en PostgreSQL)."),
    ("ADR-006", "Decimal en cálculos financieros",
     "Márgenes, IGV y totales terminan en registros contables inmutables.",
     "decimal.Decimal (ROUND_HALF_UP) en toda la aritmética financiera, nunca float.",
     "Totales exactos y reproducibles. Costo: se convierte a float solo en el borde de serialización JSON."),
    ("ADR-007", "Audit log inmutable en vez de soft delete",
     "validation_logs es el rastro de compliance ante auditoría.",
     "Tabla append-only: sin UPDATE ni DELETE en producción.",
     "Historial completo verificable. Costo: la tabla crece sin límite (volumen bajo; archivado futuro si crece)."),
    ("ADR-008", "JSON Schema Draft-07 + Pydantic (doble validación)",
     "Cada MCPEnvelope debe validarse en productor y consumidor, y también desde el frontend.",
     "Pydantic v2 para runtime en Python + JSON Schema Draft-07 validado en cada hop.",
     "El mismo contrato valida desde TypeScript/tests en cualquier lenguaje. Costo: contrato definido dos veces (riesgo de divergencia, mitigado con agent_contracts.md)."),
    ("ADR-009", "RAG con pgvector + embeddings Ollama, sin retrievers LangChain",
     "SalesAgent necesitaba recuperación semántica; ItineraryAgent usaba un diccionario estático. PostgreSQL ya es el almacén único.",
     "pgvector para embeddings de paquetes y destination_knowledge, generados con nomic-embed-text (768 dim). Recuperación vía endpoints propios.",
     "RAG genuino sin almacén adicional ni reintroducir LangChain. Ante fallo de Ollama, 503 → fallback determinístico."),
    ("ADR-010", "Salida estructurada forzada (constrained decoding) sobre prompt-only",
     "Antes cada agente pedía 'Return ONLY JSON' y parseaba con regex, sin garantía; con un 8B fallaba seguido.",
     "Modelo Pydantic → model_json_schema() como response_schema → Ollama 'format' (constrained decoding). parse_structured_output valida como segunda capa.",
     "Menos uso del fallback cuando Ollama está disponible; destapó y corrigió código muerto (_estimate_with_swarms) y la promesa sin implementar del confidence < 0.7 (HITL)."),
    ("ADR-011", "Evaluación local (golden set) en vez de LangSmith",
     "El plan de evaluación necesita golden set y procedimiento repetible; LangSmith requiere cuenta externa y el sistema no usa LangChain/LangGraph.",
     "Golden set local (tests/evaluation) + trazas en agent_interaction_logs + reporte imprimible (scripts/run_evaluation.py).",
     "Determinístico y ejecutable en CI sin infraestructura externa. La evaluación end-to-end contra el LLM real sigue siendo manual."),
    ("ADR-012", "Agentes especializados de flujo fijo en vez de patrón Deep Agent",
     "El patrón Deep Agent maneja incertidumbre sobre qué pasos son necesarios; aquí los pasos están definidos por reglas de negocio y compliance.",
     "9 agentes especializados con responsabilidad fija coordinados por una Saga explícita.",
     "Cada paso es rastreable 1:1 a un agente (auditable). Costo: no maneja tareas verdaderamente abiertas sin añadir lógica/tool explícita."),
]
for aid, title_, ctx, dec, cons in _adrs:
    h3(aid + " — " + title_)
    para("Contexto: " + ctx)
    para("Decisión: " + dec)
    para("Consecuencias: " + cons)


# ══════════════════════════════════════════════════════════════════════════════
# 5. PLAN DE EVALUACIÓN
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Plan de evaluación")

h2("5.1 Conjunto de evaluación (golden set)")
para("tests/evaluation/golden_set.py — 12 casos sobre los 4 agentes que invocan LLM (Sales, "
     "Quotation, Itinerary, Orchestrator/Fase 3), cada uno con: una salida de LLM representativa, "
     "el modelo Pydantic contra el que debe validar, si se espera válida o no, y campos a "
     "verificar. Cubre tres categorías por agente: JSON limpio, JSON válido envuelto en prosa, y "
     "salidas inválidas.")

h2("5.2 Métricas")
table(
    ["Métrica", "Fuente", "Qué mide"],
    [
        ["% casos golden set que pasan", "scripts/run_evaluation.py", "Robustez de la validación estructurada"],
        ["et_llm_call_duration_seconds (p50/p95)", "Prometheus", "Latencia de llamadas LLM por agente"],
        ["et_llm_tokens_total", "Prometheus", "Consumo de tokens por agente/modelo"],
        ["agent_interaction_logs.success", "PostgreSQL", "Tasa de éxito real de interacciones LLM"],
        ["et_agent_errors_total", "Prometheus", "Tasa de error por agente"],
    ],
)

h2("5.3 LangSmith — observabilidad y evaluación")
para("Decisión: NO se usa LangSmith (ADR-011). El proyecto no tiene cuenta y el sistema no usa "
     "LangChain/LangGraph en su camino principal, así que el tracing solo cubriría una fracción "
     "del flujo. Se documentan los equivalentes locales.")
h3("5.3.1 Tracing en desarrollo y producción")
para("Cada llamada LLM se reporta a agent_interaction_logs vía "
     "base_agent.report_llm_interaction() → POST /api/v1/agent-interactions, con input, output, "
     "duración y éxito/error.")
h3("5.3.2 Datasets y evaluadores")
para("El golden set cumple el rol de dataset versionado; parse_structured_output() + los asserts "
     "de test_golden_set.py cumplen el rol de evaluador.")
h3("5.3.3 Comparación de experimentos")
para("No automatizada: comparar corridas requiere ejecutar 'python scripts/run_evaluation.py "
     "--json' en cada punto del tiempo y diffear. Mejora futura si el proyecto escala.")
h3("5.3.4 Evaluación en línea (online evals)")
para("No implementada: requeriría un evaluador LLM-as-judge sobre tráfico real. Con Ollama local "
     "el costo marginal es bajo, así que es la extensión más natural para evaluación continua.")
h2("5.4 Procedimiento")
code("# Evaluación estructural (determinística, sin Ollama, corre en CI)\n"
     "python scripts/run_evaluation.py\n\n"
     "# Evaluación end-to-end contra el sistema real (requiere Ollama + docker compose up)\n"
     "python scripts/demo_flow.py --scenario ALL\n\n"
     "# Consultar trazas de interacciones LLM reales tras la demo\n"
     "docker compose exec postgres psql -U etuser -d everywheretravel \\\n"
     "  -c \"SELECT agent_id, action, success, duration_ms, tokens_used \\\n"
     "      FROM agent_interaction_logs ORDER BY created_at DESC LIMIT 20;\"")
h2("5.5 Reporte de resultados")
para("scripts/run_evaluation.py imprime un reporte tabular (PASS/FAIL por caso + agregado por "
     "agente) y admite --json para CI. Última corrida (17-07-2026): 12/12 casos (100%).")
table(
    ["ID", "Agente", "Resultado", "Descripción"],
    [
        ["sales-01", "sales-agent", "PASS", "JSON limpio válido"],
        ["sales-02", "sales-agent", "PASS", "JSON válido envuelto en bloque ```json``` + prosa"],
        ["sales-03", "sales-agent", "PASS", "Faltan campos requeridos (traveler_count, budget_range)"],
        ["quotation-01", "quotation-agent", "PASS", "JSON limpio válido"],
        ["quotation-02", "quotation-agent", "PASS", "JSON válido con prosa antes/después"],
        ["quotation-03", "quotation-agent", "PASS", "Respuesta sin JSON (rehúsa la tarea)"],
        ["itinerary-01", "itinerary-agent", "PASS", "JSON limpio válido con 1 día"],
        ["itinerary-02", "itinerary-agent", "PASS", "Faltan campos requeridos (days, included_services, ...)"],
        ["conflict-val-01", "orchestrator-agent", "PASS", "Validación con alta confianza"],
        ["conflict-val-02", "orchestrator-agent", "PASS", "Validación con baja confianza (< 0.7, debe escalar)"],
        ["conflict-mon-01", "orchestrator-agent", "PASS", "Monitoring válido"],
        ["conflict-mon-02", "orchestrator-agent", "PASS", "Tipos incorrectos en la salida"],
    ],
)
para("Total: 12/12 casos pasaron (100%). Por agente: sales-agent 3/3, quotation-agent 3/3, "
     "itinerary-agent 2/2, orchestrator-agent 4/4.")


# ══════════════════════════════════════════════════════════════════════════════
# 6. CATÁLOGO DE PROMPTS
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Catálogo de prompts")
para("Todos los prompts viven como texto plano versionado en el repositorio "
     "(agents/<agente>/prompts/system_prompt.txt). De los 10 agentes, solo 4 invocan realmente un "
     "LLM (Sales, Quotation, Itinerary, Orchestrator); los otros 6 cargan su prompt como "
     "documentación de rol pero su lógica es determinística. La salida de cada agente LLM se "
     "fuerza a un JSON Schema (ver sección 3.7).")

h2("Texto completo de los prompts")
_prompts = [
    ("SalesAgent", "agents/sales/prompts/system_prompt.txt"),
    ("QuotationAgent", "agents/quotation/prompts/system_prompt.txt"),
    ("ItineraryAgent", "agents/itinerary/prompts/system_prompt.txt"),
    ("OrchestratorAgent", "agents/orchestrator/prompts/system_prompt.txt"),
    ("ValidationAgent", "agents/validation/prompts/system_prompt.txt"),
    ("ReservationAgent", "agents/reservation/prompts/system_prompt.txt"),
    ("FinanceAgent", "agents/finance/prompts/system_prompt.txt"),
    ("DocumentAgent", "agents/document/prompts/system_prompt.txt"),
    ("MonitoringAgent", "agents/monitoring/prompts/system_prompt.txt"),
    ("NotificationAgent", "agents/notification/prompts/system_prompt.txt"),
]
for agent_name, rel in _prompts:
    h3(agent_name + " — " + rel)
    code(read_prompt(rel))

h3("Orchestrator Fase 3 — prompts inline (agents/orchestrator/agent.py)")
code("# _VALIDATION_AGENT_PROMPT (salida forzada: ConflictValidationOutput)\n"
     "You are a Validation Agent for conflict scenarios.\n"
     "Assess whether the conflicting state is a data integrity issue.\n"
     "Respond with JSON: { \"is_integrity_issue\": bool, \"confidence\": float, \"recommendation\": str }\n\n"
     "# _MONITORING_AGENT_PROMPT (salida forzada: ConflictMonitoringOutput)\n"
     "You are a Monitoring Agent for conflict scenarios.\n"
     "Assess the operational impact and decide if escalation to human is needed.\n"
     "Respond with JSON: { \"needs_escalation\": bool, \"impact\": \"low|medium|high|critical\", \"action\": str }")


# ══════════════════════════════════════════════════════════════════════════════
# 7. MEDICIÓN DE ÉXITO Y ROI
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Medición de éxito y ROI")
para("Nota metodológica: Everywhere Travel es una agencia ficticia para efectos del proyecto "
     "académico. Las cifras son estimaciones ilustrativas basadas en supuestos explícitos, no "
     "datos reales de negocio; el objetivo es demostrar la metodología de cálculo de ROI.")

h2("7.1 KPIs de negocio")
table(
    ["KPI", "Línea base (manual)", "Objetivo con el sistema"],
    [
        ["Tiempo de cotización", "2-4 horas", "< 30s (catálogo) / minutos (personalizado)"],
        ["Errores de doble reserva", "Frecuentes, no cuantificados", "0 (optimistic locking)"],
        ["Emisión de documentos", "Manual en Word, minutos-horas", "Automático, segundos (cola async ×3)"],
        ["Errores de cálculo financiero", "Riesgo real con Excel", "0 (aritmética Decimal)"],
        ["Disponibilidad ante fallo parcial", "Un ausente detiene el proceso", "Circuit breaker + dead-letter + escalación"],
    ],
)

h2("7.2 Línea base (baseline pre-IA)")
para("Un agente de ventas dedica 2-4 horas a armar una cotización de paquete personalizado "
     "(buscar disponibilidad, calcular precio con margen e IGV a mano, redactar itinerario). Con "
     "el sistema, esa tarea se resuelve en segundos para catálogo y en minutos para "
     "personalizados (limitado por la latencia del LLM local).")

h2("7.3 Cálculo de ROI")
para("Todas las cifras están en soles (S/). Son estimaciones ilustrativas con supuestos "
     "explícitos —no datos reales de la operación—; el objetivo es demostrar la metodología de "
     "cálculo.")

h3("Supuestos del modelo (caso base)")
table(
    ["Parámetro", "Valor", "Justificación"],
    [
        ["Costo laboral cargado del vendedor", "S/ 20 / hora",
         "Sueldo base + beneficios sociales peruanos (~45%: EsSalud, CTS, gratificaciones, vacaciones)"],
        ["Volumen de cotizaciones", "200 / mes (~50 / semana)", "Agencia con varias sedes"],
        ["Mix catálogo / personalizado", "60% / 40%", "Los personalizados son los más costosos de armar a mano"],
        ["Tiempo manual (AS-IS)", "Catálogo 1.0 h · Personalizado 3.0 h", "Rango observado 2-4 h para personalizados"],
        ["Tiempo con el sistema (TO-BE)", "Catálogo ~5 min · Personalizado ~20 min", "Incluye la revisión humana del resultado"],
        ["Ahorro por cotización", "Catálogo 0.92 h · Personalizado 2.67 h", "Diferencia AS-IS − TO-BE"],
        ["Dobles reservas evitadas", "2 / mes × S/ 300", "Reacomodo, pérdida de margen y goodwill por incidente"],
    ],
)

h3("Costos del proyecto")
para("Inversión inicial (una sola vez) — desarrollo valorizado a precio de mercado. En el "
     "proyecto académico no se factura, pero se incluye para un cálculo de ROI realista:")
table(
    ["Concepto", "Cálculo", "Monto"],
    [["Desarrollo del sistema (una vez)", "1 desarrollador × 3 meses × S/ 6,000/mes", "S/ 18,000"]],
)
para("Costos recurrentes (mensuales):")
table(
    ["Concepto", "Cálculo", "Monto / mes"],
    [
        ["Infraestructura (VPS gama media con RAM para Ollama + servicios)", "Alquiler mensual del servidor", "S/ 300"],
        ["LLM (Ollama local, qwen3:8b)", "Costo por token", "S/ 0"],
        ["Mantenimiento (monitoreo, actualizaciones)", "6 h/mes × S/ 40/h", "S/ 240"],
        ["Total recurrente", "", "S/ 540"],
    ],
)

h3("Beneficios cuantificables")
para("El principal es el ahorro de horas-hombre en cotización; se suma la evitación de costos por "
     "doble reserva:")
table(
    ["Beneficio", "Cálculo (mensual)", "Monto / mes"],
    [
        ["Ahorro laboral — cotizaciones de catálogo", "120 × 0.92 h × S/ 20", "S/ 2,208"],
        ["Ahorro laboral — cotizaciones personalizadas", "80 × 2.67 h × S/ 20", "S/ 4,272"],
        ["Doble reserva evitada", "2 × S/ 300", "S/ 600"],
        ["Beneficio bruto", "", "S/ 7,080"],
    ],
)
para("No se monetizan aquí otros beneficios reales pero difíciles de cuantificar: eliminación de "
     "errores de cálculo financiero (Decimal exacto), reducción de reprocesos por versiones de "
     "cotización, y emisión de documentos sin errores de tipeo. El ROI calculado es, por tanto, "
     "conservador.")

h3("Fórmula de ROI y cálculo (caso base)")
code("ROI = (Beneficio neto acumulado − Costo total) / Costo total × 100\n\n"
     "Beneficio neto mensual  = 7,080 − 540               = S/ 6,540 / mes\n"
     "Payback (recuperación)  = 18,000 / 6,540            ~= 2.8 meses\n\n"
     "Horizonte 12 meses:\n"
     "  Beneficio bruto anual = 7,080 x 12                = S/ 84,960\n"
     "  Costo total ano 1     = 18,000 + (540 x 12)       = S/ 24,480\n"
     "  Beneficio neto anual  = 84,960 - 24,480           = S/ 60,480\n"
     "  ROI (12 meses)        = 60,480 / 24,480 x 100     ~= 247%")

h3("Análisis de sensibilidad")
para("El resultado depende sobre todo del volumen de cotizaciones y del costo laboral. Tres "
     "escenarios:")
table(
    ["Escenario", "Supuestos", "Beneficio neto/mes", "Payback", "ROI 12 meses"],
    [
        ["Conservador", "120 cot/mes · S/ 18/h · 1 doble reserva", "S/ 3,259", "~5.5 meses", "~86%"],
        ["Base", "200 cot/mes · S/ 20/h · 2 dobles reservas", "S/ 6,540", "~2.8 meses", "~247%"],
        ["Optimista", "320 cot/mes · S/ 22/h · 3 dobles reservas", "S/ 11,765", "~1.5 meses", "~503%"],
    ],
)
para("Incluso en el escenario conservador el proyecto se recupera en menos de 7 meses y arroja un "
     "ROI positivo en el primer año, lo que sostiene la decisión de automatizar. El caso base "
     "indica una recuperación de la inversión en menos de 3 meses.")

h2("7.4 Tablero de éxito (técnico + negocio)")
table(
    ["Dimensión", "Métrica", "Dónde se mide"],
    [
        ["Técnico", "Latencia p95, tasa de error por agente, estado de circuit breakers", "Grafana 'Everywhere Travel — Overview'"],
        ["Técnico", "% golden set que pasa", "scripts/run_evaluation.py (CI en cada PR)"],
        ["Negocio", "Tiempo promedio de cotización a VALIDATED", "sagas (created_at → paso validation_complete)"],
        ["Negocio", "Tasa de sagas COMPLETED sin intervención manual", "sagas.status"],
        ["Negocio", "Tasa de escalación a humano (HITL)", "ConflictResolved.needs_escalation / logs de Monitoring"],
    ],
)

h2("7.5 Cadencia de revisión")
bullets([
    ("Diaria (automática)", "Alertas Prometheus: circuit breakers, DLQ, latencia LLM."),
    ("Semanal", "Dashboard Grafana + agent_interaction_logs para detectar degradación de calidad del LLM."),
    ("Por release/PR", "Golden set en CI: ningún cambio se mergea si falla."),
    ("Mensual (con datos reales)", "Recalcular el ROI con cifras reales de volumen y tiempo ahorrado."),
])


# ══════════════════════════════════════════════════════════════════════════════
# 8. DESPLIEGUE Y OPERACIÓN
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Despliegue y operación")

h2("8.1 Entornos")
table(
    ["Entorno", "Existe", "Configuración", "Diferencias clave"],
    [
        ["Desarrollo", "Sí", ".env + docker compose up, Dockerfile.api con --reload",
         "Logs con color (structlog), CORS a localhost, cookie JWT sin flag secure"],
        ["Producción", "Definido, no desplegado", ".env.production + Dockerfile.prod, ENVIRONMENT=production",
         "Logs en JSON, cookie JWT secure=true; exige rotar SECRET_KEY y credenciales"],
        ["Staging", "No existe", "—", "Fuera de alcance; el equivalente es correr demo_flow.py contra el stack local"],
    ],
)

h2("8.2 CI/CD y versionado")
para(".github/workflows/ci.yml ejecuta en cada push/PR: tests unitarios, adversariales, golden "
     "set, cobertura, build de la imagen de la API (Dockerfile.api) y typecheck+build del "
     "frontend. Versionado: historial de Git (git tag por release); un pipeline de imágenes "
     "etiquetadas queda como extensión (no hay registro de contenedores configurado).")

h2("8.3 Topología de despliegue")
para("docker-compose.yml orquesta 17 servicios: postgres (pgvector), redis, rabbitmq, minio, "
     "api, orchestrator, sales_worker, quotation_worker, document_worker (×3 réplicas), "
     "monitoring_worker, reservation_worker, finance_worker, validation_worker, "
     "notification_worker, itinerary_worker, prometheus, grafana, frontend. Ollama corre en el "
     "host (no en compose); los contenedores lo alcanzan vía host.docker.internal.")

h2("8.4 Configuración y secretos")
para("Variables por entorno en .env.example (dev) y .env.production. Interruptor único ENVIRONMENT "
     "entre dev y producción. Secretos en variables de entorno planas (sin vault); SECRET_KEY con "
     "valor por defecto que debe rotarse. .env/.env.production no versionados.")

h2("8.5 Estrategias de release")
table(
    ["Estrategia", "Aplicación en este sistema"],
    [
        ["Rolling update por agente", "Cada agente es un contenedor independiente; RabbitMQ retiene mensajes durante el reinicio."],
        ["Réplicas para servicios sin estado", "document_worker en 3 réplicas; el patrón se replica sin cambios de código."],
        ["Migraciones idempotentes", "ensure_schema_compatibility() aplica cambios compatibles con rolling updates."],
        ["Feature flags", "ENABLE_INLINE_QUOTATION_PIPELINE, ENABLE_LLM_ITINERARY activan código nuevo sin desplegar otra versión."],
        ["Blue-green / canary", "No implementado; el rolling update por contenedor es suficiente al volumen actual."],
    ],
)

h2("8.6 Monitoreo y alertas")
para("Prometheus scrapea api:8000/metrics y rabbitmq. Métricas custom en core/metrics.py "
     "(latencia, tokens LLM, circuit breaker state, sagas activas, DLQ). Reglas de alerta en "
     "infrastructure/prometheus/rules/alerts.yml: CircuitBreakerOpen, DeadLetterQueueGrowing, "
     "HighAgentErrorRate, SagasStuckRunning, LLMCallLatencyHigh, APIDown. Dashboard Grafana "
     "'Everywhere Travel — Overview' provisionado automáticamente. No hay Alertmanager desplegado "
     "(las alertas se ven en la UI de Prometheus, no se enrutan a Slack/email).")

h2("8.7 Procedimiento ante incidentes")
para("Runbook con tres niveles de severidad:")
table(
    ["Severidad", "Ejemplo", "Respuesta"],
    [
        ["Crítica", "CircuitBreakerOpen, APIDown", "Revisar de inmediato — el sistema no puede vender/reservar"],
        ["Alta", "HighAgentErrorRate, DeadLetterQueueGrowing", "Revisar en el día — degradación parcial con fallbacks"],
        ["Media", "LLMCallLatencyHigh, SagasStuckRunning", "Revisar en la semana — impacto en experiencia, no en datos"],
    ],
)
para("Camino de escalación: (1) alerta Prometheus → Grafana; (2) auto-recuperación primero "
     "(circuit breaker HALF_OPEN, requeue con backoff); (3) si falla 3 veces, "
     "MonitoringAgent escala a humano (REQUIRES_MANUAL_INTERVENTION → WebSocket); (4) el operador "
     "consulta agent_interaction_logs / validation_logs / sagas (inmutables) para diagnosticar.")

h2("8.8 Escalado y FinOps")
para("Escalado horizontal presente: document_worker en 3 réplicas (la generación de PDF es lo "
     "más costoso en CPU). Cuellos de botella ante escalado: Ollama (un proceso; pool de "
     "instancias o migrar a proveedor de pago), PostgreSQL (réplicas de lectura para RAG), "
     "RabbitMQ (clustering). FinOps: con Ollama local el costo por LLM es S/ 0 (solo el hardware, "
     "ya amortizado); el costo operativo es la infraestructura (~S/ 300/mes, ver sección 7.3). Si "
     "se migrara a un proveedor de pago, et_llm_tokens_total sería la base para proyectar el costo "
     "(tokens × precio_por_token, agrupable por agente).")


# ══════════════════════════════════════════════════════════════════════════════
# 9. APÉNDICES
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Apéndices")

h2("9.1 Glosario")
table(
    ["Término", "Definición"],
    [
        ["Agente", "Proceso especializado que ejecuta una responsabilidad del negocio y se comunica vía mensajes MCP."],
        ["Saga", "Patrón de transacción distribuida: secuencia de pasos con compensación ante fallos, sin bloqueo global."],
        ["MCP Envelope", "Contrato de mensajería inter-agente (message_id, saga_id, sender/receiver, payload_type, payload, TTL, prioridad)."],
        ["RAG", "Retrieval-Augmented Generation: recuperación de conocimiento por similaridad semántica para enriquecer al LLM."],
        ["pgvector", "Extensión de PostgreSQL para almacenar y consultar vectores de embedding (distancia de coseno)."],
        ["Embedding", "Vector numérico (aquí 768 dim) que representa el significado de un texto para búsqueda semántica."],
        ["Constrained decoding", "Restricción del LLM durante la generación para que la salida cumpla un JSON Schema."],
        ["Circuit breaker", "Patrón que corta llamadas a un servicio degradado (OPEN) y prueba su recuperación (HALF_OPEN)."],
        ["Dead-letter queue (DLQ)", "Cola donde caen los mensajes que no se pudieron procesar, para reintento o inspección."],
        ["HITL (Human-in-the-loop)", "Punto donde el sistema escala una decisión a un operador humano."],
        ["IGV", "Impuesto General a las Ventas del Perú (18%)."],
        ["Golden set", "Conjunto curado de casos de evaluación con salida esperada, usado como referencia de calidad."],
        ["Optimistic locking", "Bloqueo no bloqueante (aquí Redis SETNX) que asume que los conflictos son raros."],
        ["Ollama", "Runtime local de modelos LLM; aquí sirve qwen3:8b y nomic-embed-text."],
        ["ADR", "Architecture Decision Record: registro inmutable de una decisión de arquitectura."],
    ],
)

h2("9.2 Referencias")
h3("Referencias bibliográficas (formato APA 7)")
para("Fundamentos académicos:")
bullets([
    "García-Molina, H., & Salem, K. (1987). Sagas. En Proceedings of the 1987 ACM SIGMOD "
    "International Conference on Management of Data (pp. 249–259). ACM. "
    "https://doi.org/10.1145/38713.38742",
    "Lewis, P., Perez, E., Piktus, A., et al. (2020). Retrieval-augmented generation for "
    "knowledge-intensive NLP tasks. Advances in Neural Information Processing Systems 33 "
    "(NeurIPS 2020). https://arxiv.org/abs/2005.11401",
    "Guo, T., et al. (2024). Large language model based multi-agents: A survey of progress and "
    "challenges. Proceedings of IJCAI 2024 (pp. 8048–8057). https://arxiv.org/abs/2402.01680",
])
para("Patrones de arquitectura:")
bullets([
    "Fowler, M. (2014, 6 de marzo). Circuit Breaker. martinfowler.com. "
    "https://martinfowler.com/bliki/CircuitBreaker.html",
    "Nygard, M. T. (2018). Release It! Design and Deploy Production-Ready Software (2.ª ed.). "
    "Pragmatic Bookshelf.",
    "MADR. (s. f.). Markdown Architecture Decision Records. https://adr.github.io/madr/",
])
para("Estándares y documentación técnica oficial:")
bullets([
    "Anthropic. (2024, 25 de noviembre). Introducing the Model Context Protocol. "
    "https://www.anthropic.com/news/model-context-protocol · https://modelcontextprotocol.io",
    "Ollama. (2024, 6 de diciembre). Structured outputs. https://ollama.com/blog/structured-outputs "
    "· https://docs.ollama.com/capabilities/structured-outputs",
    "Object Management Group. (2011). Business Process Model and Notation (BPMN) Version 2.0. "
    "https://www.omg.org/spec/BPMN/2.0/",
    "pgvector. Open-source vector similarity search for Postgres. "
    "https://github.com/pgvector/pgvector",
    "Pydantic. Documentation. https://docs.pydantic.dev",
    "FastAPI. Documentation. https://fastapi.tiangolo.com",
    "RabbitMQ. Documentation. https://www.rabbitmq.com/docs",
    "Redis. Documentation. https://redis.io/docs",
    "Prometheus. Documentation. https://prometheus.io/docs",
])
h3("Documentación interna del proyecto")
bullets([
    "README.md — visión general, agentes, stack, despliegue.",
    "docs/architecture.md — arquitectura, diagramas de secuencia/estado/ER, RAG, HITL.",
    "docs/adr/ — los 12 ADR completos (contexto, alternativas, consecuencias).",
    "docs/proceso_negocio.md y docs/bpmn/ — narrativa y diagramas BPMN 2.0 AS-IS/TO-BE.",
    "docs/tools_catalog.md, docs/prompts_catalog.md — fichas de tools y texto de prompts.",
    "docs/security.md, docs/evaluation.md, docs/roi.md — seguridad, evaluación y ROI.",
])


# ─── Guardar ──────────────────────────────────────────────────────────────────
out = ROOT / "docs" / "informe_final.docx"
doc.save(str(out))
print("Documento generado:", out)
